import sys
import os
import time
import pathlib
import shutil
import ruamel.std.zipfile as zipfile

from docx import Document
from os import rmdir, removedirs
from os.path import join, expanduser
from bs4 import BeautifulSoup


_root_ = str(pathlib.Path(__file__).parent.absolute())
_file_ = ""
docx = ""

image = {
    "name": "",
    "path": "",
}

skip_file = [
    "[Content_Types].xml"
]

should_extract = {"files": "", "paths": ""}


def ext_img_from_docx(docx_file):
    """extract the image that contain the content of docx file"""

    archive = zipfile.ZipFile(f"{_root_}\\docx\\{docx_file}")
    for file in archive.filelist:
        if file.filename.startswith("word/media/"):
            image["name"] = file.filename.split("/")[-1]
            image["path"] = "%s\\archive\\%s" % (_root_, image["name"])
            archive.extract(file.filename, path=f"{_root_}/archive/")

    # current path of image
    current_path = "%s\\archive\\word\\media\\%s" % (_root_, image["name"])
    # move image to `archive dir
    pathlib.Path(current_path).rename(image["path"])
    # remore word directory
    shutil.rmtree(_root_ + "\\archive\\word")


def ext_docx_content_from_image():
    """Extract docx content from image"""

    global should_extract
    files = []
    paths = []

    try:
        with zipfile.ZipFile(image["path"]) as content:
            # extract all docx content from image
            content.extractall(path="archive\\content")

            for file in content.filelist:
                this_file = file.filename.split("/")[-1]

                # check if extracted file only extracted if file is not
                # in skipped file lists
                if this_file not in skip_file:
                    files.append(this_file)
                    paths.append(file.filename)

            # append info of files and paths
            should_extract["files"] = files
            should_extract["paths"] = paths

            # create new dir
            pathlib.Path(_root_ + "\\archive\\docx_content\\").mkdir(
                parents=True, exist_ok=True
            )

            count = 0
            for file in should_extract["paths"]:
                current_path = "%s\\archive\\content\\%s" % (
                    _root_,
                    should_extract["paths"][count],
                )

                # move all extracted files to `docx_content`
                pathlib.Path(current_path).rename(
                    "%s%s"
                    % (
                        _root_ + "\\archive\\docx_content\\",
                        should_extract["files"][count],
                    )
                )
                count += 1

            # delete dir
            shutil.rmtree(_root_ + "\\archive\\content")
    except Exception as ex:
        print(ex)


def create_docx_template(name="result.docx"):
    """Create docx file template for merging docx"""
    global docx

    docx = Document()
    docx.add_heading("Template Document", 0)
    docx.save(name)


def merging_docx():
    """Merging content of file docx with template docx"""

    mime_type = (".jpg", ".png", ".jpeg")

    # delete file that have same name as the file that should be extracted
    print(should_extract)
    index = 0
    for file in should_extract["paths"]:
        zipfile.delete_from_zip_file(
            "result.docx", pattern=should_extract["paths"][index]
        )

        index += 1

    # python run.py -f tugas_img_test.docx
    with zipfile.ZipFile("result.docx", "a") as content:
        try:
            index = 0
            for file in should_extract["paths"]:
                file_index_path = should_extract["paths"][index]
                file_path = "%s\\archive\\docx_content\\%s" % (
                    _root_,
                    should_extract["files"][index],
                )

                with open(file_path, "r", encoding="utf-8") as f:
                    if file_index_path.endswith(mime_type):
                        image_path = "%s%s" % (_root_ + "\\archive\\docx_content\\", should_extract["files"][index])

                        image_handle = open(image_path, "rb")
                        raw_image_data = image_handle.read()
                        content.writestr(
                            file_index_path,
                            raw_image_data
                        )
                        image_handle.close()

                        # docx.add_picture(
                        #     "%s%s"
                        #     % (
                        #         _root_ + "\\archive\\docx_content\\",
                        #         should_extract["files"][index],
                        #     )
                        # )

                    else:
                        data = f.read()

                        # get content data of file in `docx_content` dir
                        content_data = BeautifulSoup(data, "xml")
                        # overwrite the files content with content in `docx_content` dir
                        content.writestr(
                            file_index_path,
                            bytes(str(content_data), encoding="utf-8"),
                        )
                        f.close()

                index += 1

            # delete dir
            shutil.rmtree(_root_ + "\\archive")

            print("File has been merged.")
            print("Completed.")
        except Exception as ex:
            print(ex)


def command(args):
    """
    :Command Helper

    --file -f file name that will be extracted; `C://user/document/file.docx`
    --name -n file name that will be save; default `result.docx`
    """
    try:
        file = args[1]

        data = {"file": ""}
        if args[0] == ("-f" or "--file"):
            if file.endswith(".docx"):
                data["file"] = file

        return data
    except Exception as ex:
        print(err, "\n", str(ex.args))


def main():
    start_time = time.time()
    global _file_, _name_
    response = command(args=sys.argv[1:])

    _file_ = response["file"]

    ext_img_from_docx(docx_file=_file_)
    create_docx_template()
    ext_docx_content_from_image()
    merging_docx()
    print("--- %s seconds ---" % (time.time() - start_time))


if __name__ == "__main__":
    main()